#!/usr/bin/env python3
"""Genera 4 versioni della verifica di Sistemi e Reti in formato DOCX."""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

OUTPUT_DIR = "/home/user/presentazione-vpn/verifiche"
os.makedirs(OUTPUT_DIR, exist_ok=True)

HEADER_COMMON = (
    "VERIFICA DI SISTEMI E RETI",
    "Firewall, ACL, DMZ, NAT e VPN",
)

FOOTER = (
    "Durata massima della prova: 3 ore.\n"
    "È consentito l'uso di manuali tecnici (reference con solo sintassi, non guide).\n"
    "Non è consentito lasciare l'aula prima che sia trascorsa 1 ora e 30 minuti dall'inizio della prova."
)

INSTRUCTIONS = (
    "Il candidato svolga la prima parte della prova e risponda a uno dei quesiti "
    "proposti nella seconda parte."
)

verifiche = [
    {
        "versione": "A",
        "titolo_scenario": "MediClinic",
        "scenario": (
            "Una clinica medica privata ha una sede principale a Bologna con 20 postazioni "
            "distribuite tra ambulatori, reception e uffici amministrativi, e 3 ambulatori "
            "periferici in città vicine, ciascuno con 5 postazioni. La sede principale gestisce "
            "un archivio di cartelle cliniche elettroniche, un servizio web per la prenotazione "
            "online delle visite accessibile al pubblico e un servizio di posta elettronica. "
            "I medici degli ambulatori periferici devono poter consultare e aggiornare le "
            "cartelle cliniche in tempo reale. I dati trattati sono soggetti a normative "
            "stringenti sulla privacy."
        ),
        "prima_parte": [
            (
                "Proponga un progetto, anche grafico, dell'infrastruttura di rete della sede "
                "principale, motivando le scelte relative alla topologia, ai dispositivi di rete, "
                "alla segmentazione della rete e al posizionamento dei servizi esposti al pubblico "
                "rispetto a quelli interni. Indichi le misure adottate per la protezione della "
                "rete e dei dati."
            ),
            (
                "Definisca le politiche di filtraggio del traffico tra le diverse zone della rete, "
                "specificando quali flussi di comunicazione devono essere permessi o negati. "
                "Illustri la strategia adottata per la gestione degli indirizzi IP e la "
                "pubblicazione dei servizi verso l'esterno."
            ),
            (
                "Illustri la soluzione adottata per consentire agli ambulatori periferici di "
                "accedere in modo sicuro ai servizi della sede principale attraverso la rete "
                "Internet, motivando la scelta dei protocolli e delle tecnologie utilizzate, "
                "con particolare attenzione alla protezione dei dati sanitari."
            ),
        ],
        "seconda_parte": [
            (
                "Si descrivano le diverse tipologie di firewall, illustrando per ciascuna il "
                "livello di analisi del traffico, i vantaggi e gli svantaggi, e indicando in "
                "quali contesti è preferibile l'una o l'altra."
            ),
            (
                "Si illustrino i meccanismi attraverso cui una VPN garantisce riservatezza, "
                "integrità e autenticazione delle comunicazioni, descrivendo almeno due protocolli "
                "utilizzabili e le relative differenze."
            ),
        ],
    },
    {
        "versione": "B",
        "titolo_scenario": "EcoTour",
        "scenario": (
            "La società EcoTour gestisce una rete di 5 agriturismi distribuiti in Toscana. "
            "La sede amministrativa si trova a Firenze con 10 postazioni e gestisce un portale "
            "web per le prenotazioni online, un servizio di posta elettronica e un database "
            "centralizzato con disponibilità, tariffe e dati dei clienti. Ogni agriturismo "
            "dispone di 3 postazioni per il personale e offre connettività wireless agli ospiti. "
            "Tre responsabili di zona si spostano frequentemente tra le strutture e necessitano "
            "di accesso al gestionale aziendale da dispositivi mobili in qualsiasi momento."
        ),
        "prima_parte": [
            (
                "Proponga un progetto, anche grafico, dell'infrastruttura di rete della sede di "
                "Firenze, motivando le scelte relative alla topologia, ai dispositivi, alla "
                "segmentazione della rete e alla separazione tra i servizi accessibili dall'esterno "
                "e quelli interni. Discuta inoltre come garantire negli agriturismi la separazione "
                "tra la rete del personale e quella degli ospiti."
            ),
            (
                "Definisca le politiche di filtraggio del traffico della sede principale, "
                "specificando le regole di comunicazione tra le diverse zone della rete. Illustri "
                "la strategia adottata per la gestione degli indirizzi IP e per rendere i servizi "
                "pubblici raggiungibili dall'esterno."
            ),
            (
                "Proponga le soluzioni per consentire il collegamento sicuro degli agriturismi "
                "alla sede centrale e l'accesso remoto dei responsabili di zona. Confronti almeno "
                "due possibili approcci, motivando la scelta finale."
            ),
        ],
        "seconda_parte": [
            (
                "Si descriva il funzionamento delle tecniche di traduzione degli indirizzi di "
                "rete, illustrando con esempi le diverse modalità e i relativi casi d'uso."
            ),
            (
                "Si confrontino almeno due diverse tecnologie per la realizzazione di reti "
                "private virtuali, evidenziando differenze architetturali, vantaggi e svantaggi "
                "di ciascun approccio."
            ),
        ],
    },
    {
        "versione": "C",
        "titolo_scenario": "DigitalPress",
        "scenario": (
            "Una casa editrice digitale ha la sede principale a Milano con 25 postazioni "
            "distribuite tra redazione, reparto grafico e amministrazione, e una sede secondaria "
            "a Roma con 10 postazioni dedicate alla redazione. La sede di Milano ospita il sito "
            "web pubblico del catalogo, una piattaforma intranet per la collaborazione tra "
            "redattori e un archivio dei contenuti digitali. I redattori di Roma devono accedere "
            "alla piattaforma intranet e all'archivio come se fossero in sede. Alcuni "
            "collaboratori freelance lavorano da casa e necessitano di accesso alla sola "
            "piattaforma intranet."
        ),
        "prima_parte": [
            (
                "Proponga un progetto, anche grafico, dell'infrastruttura di rete della sede di "
                "Milano, motivando le scelte relative alla topologia, ai dispositivi, alla "
                "segmentazione della rete e alla separazione tra servizi pubblici e risorse "
                "interne. Indichi le misure adottate per la protezione della rete da accessi "
                "non autorizzati."
            ),
            (
                "Definisca le politiche di filtraggio del traffico tra le diverse zone della rete, "
                "specificando le regole che garantiscono l'accessibilità dei servizi pubblici senza "
                "esporre le risorse interne. Illustri la strategia adottata per la gestione degli "
                "indirizzi IP."
            ),
            (
                "Illustri le soluzioni adottate per consentire alla sede di Roma e ai "
                "collaboratori freelance di accedere in modo sicuro alle risorse della sede "
                "principale, evidenziando le differenze tra i due scenari di collegamento e "
                "motivando le scelte tecnologiche."
            ),
        ],
        "seconda_parte": [
            (
                "Si descriva il concetto di zona demilitarizzata in un'infrastruttura di rete, "
                "illustrando le possibili architetture e i vantaggi in termini di sicurezza."
            ),
            (
                "Si illustri come avviene la negoziazione e l'instaurazione di un tunnel VPN "
                "sicuro tra due sedi, descrivendo i passaggi principali e i meccanismi "
                "crittografici coinvolti."
            ),
        ],
    },
    {
        "versione": "D",
        "titolo_scenario": "SmartLogistics",
        "scenario": (
            "L'azienda SmartLogistics gestisce un magazzino automatizzato a Torino (sede "
            "principale, 15 postazioni tra uffici e sala controllo) e 2 punti di distribuzione "
            "a Genova e Verona con 8 postazioni ciascuno. La sede di Torino ospita il gestionale "
            "del magazzino, un portale web tramite il quale i clienti possono monitorare lo stato "
            "delle proprie spedizioni e un servizio di posta elettronica. Nel magazzino sono "
            "presenti dispositivi automatizzati che comunicano con il gestionale sulla rete "
            "locale. I punti di distribuzione devono accedere in tempo reale al gestionale per "
            "aggiornare lo stato delle consegne."
        ),
        "prima_parte": [
            (
                "Proponga un progetto, anche grafico, dell'infrastruttura di rete della sede di "
                "Torino, motivando le scelte relative alla topologia, ai dispositivi, alla "
                "segmentazione della rete e alla separazione tra le diverse tipologie di traffico. "
                "Indichi le misure di sicurezza adottate."
            ),
            (
                "Definisca le politiche di filtraggio del traffico tra le diverse zone della rete, "
                "specificando le regole di comunicazione necessarie. Illustri la strategia adottata "
                "per la gestione degli indirizzi IP e per rendere accessibili dall'esterno i "
                "servizi destinati ai clienti."
            ),
            (
                "Proponga la soluzione per collegare in modo sicuro i punti di distribuzione alla "
                "sede centrale, motivando la scelta della tecnologia e dei protocolli. Discuta "
                "come garantire la continuità del collegamento per non interrompere l'operatività "
                "aziendale."
            ),
        ],
        "seconda_parte": [
            (
                "Si descrivano le diverse tipologie di Access Control List, spiegando i criteri "
                "di filtraggio, il posizionamento nella rete e i casi d'uso tipici di ciascuna."
            ),
            (
                "Si descrivano le diverse tipologie di VPN (site-to-site e remote access), "
                "illustrando per ciascuna le caratteristiche, i protocolli utilizzabili e gli "
                "scenari d'uso più appropriati."
            ),
        ],
    },
]


def add_horizontal_line(doc):
    """Add a thin horizontal line."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr()
    from docx.oxml.ns import qn
    pBdr = pPr.makeelement(qn('w:pBdr'), {})
    bottom = pBdr.makeelement(qn('w:bottom'), {
        qn('w:val'): 'single',
        qn('w:sz'): '4',
        qn('w:space'): '1',
        qn('w:color'): '000000',
    })
    pBdr.append(bottom)
    pPr.append(pBdr)


def create_verifica(v):
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    # -- HEADER --
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("VERIFICA DI SISTEMI E RETI")
    run.bold = True
    run.font.size = Pt(16)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Firewall, ACL, DMZ, NAT e VPN")
    run.font.size = Pt(13)
    run.italic = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Versione {v['versione']}")
    run.bold = True
    run.font.size = Pt(12)

    # Student info
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("Classe 5ª Sez. ______   Alunno ________________________________________________   Data ___/___/______")
    run.font.size = Pt(11)

    add_horizontal_line(doc)

    # Instructions
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(INSTRUCTIONS)
    run.italic = True
    run.font.size = Pt(11)

    add_horizontal_line(doc)

    # -- PRIMA PARTE --
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    run = p.add_run("PRIMA PARTE")
    run.bold = True
    run.font.size = Pt(14)

    # Scenario
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(v["scenario"])
    run.font.size = Pt(12)

    # Prompt
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    run = p.add_run("Il candidato, formulate le opportune ipotesi aggiuntive, sviluppi i seguenti punti:")
    run.italic = True
    run.font.size = Pt(12)

    # Punti prima parte
    for i, punto in enumerate(v["prima_parte"], 1):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.left_indent = Cm(0.5)
        run = p.add_run(f"{i}. ")
        run.bold = True
        run.font.size = Pt(12)
        run = p.add_run(punto)
        run.font.size = Pt(12)

    add_horizontal_line(doc)

    # -- SECONDA PARTE --
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    run = p.add_run("SECONDA PARTE")
    run.bold = True
    run.font.size = Pt(14)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(
        "Il candidato risponda a uno dei seguenti quesiti:"
    )
    run.italic = True
    run.font.size = Pt(12)

    for i, quesito in enumerate(v["seconda_parte"], 1):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.left_indent = Cm(0.5)
        run = p.add_run(f"{i}. ")
        run.bold = True
        run.font.size = Pt(12)
        run = p.add_run(quesito)
        run.font.size = Pt(12)

    add_horizontal_line(doc)

    # -- FOOTER --
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    for line in FOOTER.split('\n'):
        run = p.add_run(line + '\n')
        run.font.size = Pt(10)
        run.italic = True

    # Save
    filename = f"Verifica_SIR_Versione_{v['versione']}.docx"
    filepath = os.path.join(OUTPUT_DIR, filename)
    doc.save(filepath)
    print(f"Creato: {filepath}")


if __name__ == "__main__":
    for v in verifiche:
        create_verifica(v)
    print(f"\nTutte le verifiche salvate in: {OUTPUT_DIR}")
